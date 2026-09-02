import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Target output directory
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../..", "data", "generated_outputs"))

def generate_docx_report(filename: str, title: str, content: str, classification: str = "CONFIDENTIAL - INTERNAL INDUSTRIAL WORK") -> str:
    """
    Generates a formal, professionally styled Word document (.docx) 
    suitable for refinery approval notes, safety reports, or plant documentation.
    """
    os.makedirs(BASE_DIR, exist_ok=True)
    
    # Ensure filename ends with .docx
    safe_filename = os.path.basename(filename)
    if not safe_filename.endswith(".docx"):
        safe_filename += ".docx"
        
    file_path = os.path.join(BASE_DIR, safe_filename)
    
    doc = Document()
    
    # Set standard margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Add Classification Header Notice
    p_header = doc.add_paragraph()
    run_header = p_header.add_run(f"[{classification}]")
    run_header.font.size = Pt(9)
    run_header.font.bold = True
    run_header.font.color.rgb = RGBColor(180, 0, 0) # Dark red warning color
    
    # Document Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph() # Spacer line
    
    # Body Content (split by newlines to form clean paragraphs)
    paragraphs = content.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)
            
    # Save the document locally (Air-gapped)
    doc.save(file_path)
    return f"Success: Formal Word document generated and saved at {file_path}"